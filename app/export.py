from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


@dataclass
class MarkdownElement:
    """Represents a parsed Markdown element."""

    kind: str
    text: str
    level: int | None = None
    number: int | None = None


def parse_markdown(markdown_text: str) -> List[MarkdownElement]:
    r"""Parse a subset of Markdown into structured elements.

    Supports headings (``#``), unordered lists (``-``/``*``), ordered lists
    (``1.``, ``2.``, ...), block quotes (``>``), fenced code blocks (```),
    blank lines, and regular paragraphs.
    """

    elements: List[MarkdownElement] = []
    in_code = False
    code_lines: list[str] = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()

        # Fenced code block handling
        if re.match(r"^\s*```", line):
            if in_code:
                # closing fence -> emit code block
                elements.append(MarkdownElement("codeblock", "\n".join(code_lines)))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(raw_line)  # preserve indentation and spacing
            continue
        if not line:
            elements.append(MarkdownElement("blank", ""))
            continue

        if m := re.match(r"^(#{1,6})\s+(.*)", line):
            elements.append(MarkdownElement("heading", m.group(2).strip(), level=len(m.group(1))))
            continue

        if m := re.match(r"^\s*[-*]\s+(.*)", line):
            elements.append(MarkdownElement("bullet", m.group(1).strip()))
            continue

        if m := re.match(r"^(\d+)\.\s+(.*)", line):
            elements.append(MarkdownElement("number", m.group(2).strip(), number=int(m.group(1))))
            continue

        if m := re.match(r"^\s*>\s?(.*)", line):
            elements.append(MarkdownElement("quote", m.group(1).strip()))
            continue

        elements.append(MarkdownElement("paragraph", line))

    # If file ended while in a code block, flush it
    if in_code:
        elements.append(MarkdownElement("codeblock", "\n".join(code_lines)))

    return elements


def _add_md_inline_runs(paragraph, text: str) -> None:
    """Add runs to a python-docx paragraph by parsing simple inline Markdown.

    Supports code spans (`code`), bold (**text**), italic (*text*). Nested
    combinations are not fully supported; code has highest precedence.
    """

    def _emit_code(p, t: str) -> None:
        run = p.add_run(t)
        run.font.name = "Consolas"
        run.font.size = Pt(11)
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

    def _emit_text_with_styles(p, t: str) -> None:
        idx = 0
        while idx < len(t):
            # Bold
            if t.startswith("**", idx):
                end = t.find("**", idx + 2)
                if end != -1:
                    run = p.add_run(t[idx + 2 : end])
                    run.bold = True
                    run.font.size = Pt(11)
                    idx = end + 2
                    continue
            # Italic (single *)
            if t[idx] == "*":
                end = t.find("*", idx + 1)
                if end != -1:
                    run = p.add_run(t[idx + 1 : end])
                    run.italic = True
                    run.font.size = Pt(11)
                    idx = end + 1
                    continue
            # Plain char
            # Group consecutive non-markup characters into one run
            next_marks = [x for x in [t.find("**", idx), t.find("*", idx)] if x != -1 and x >= idx]
            next_idx = min(next_marks) if next_marks else len(t)
            run = paragraph.add_run(t[idx:next_idx])
            run.font.size = Pt(11)
            idx = next_idx

    # Tokenize by code spans first
    i = 0
    while i < len(text):
        if text[i] == "`":
            j = text.find("`", i + 1)
            if j != -1:
                _emit_code(paragraph, text[i + 1 : j])
                i = j + 1
                continue
        # Emit non-code until next backtick or end
        j = text.find("`", i)
        segment = text[i : (j if j != -1 else len(text))]
        _emit_text_with_styles(paragraph, segment)
        if j == -1:
            break
        i = j


def export_docx(title: str, markdown_text: str) -> Tuple[str, bytes]:
    doc = Document()
    doc.add_heading(title or "特許明細書草案", level=1)
    for element in parse_markdown(markdown_text):
        if element.kind == "blank":
            doc.add_paragraph()
            continue
        if element.kind == "heading":
            level = element.level if element.level is not None else 1
            doc.add_heading(element.text, level=level)
            continue
        if element.kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_md_inline_runs(p, element.text)
            continue
        if element.kind == "number":
            p = doc.add_paragraph(style="List Number")
            _add_md_inline_runs(p, element.text)
            continue
        if element.kind == "quote":
            p = doc.add_paragraph(style="Intense Quote")
            _add_md_inline_runs(p, element.text)
            continue
        if element.kind == "codeblock":
            for line in (element.text or "").splitlines() or [""]:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            continue
        p = doc.add_paragraph()
        _add_md_inline_runs(p, element.text)
    buffer = BytesIO()
    doc.save(buffer)
    return (f"{title or 'draft'}.docx", buffer.getvalue())


def export_pdf(title: str, markdown_text: str) -> Tuple[str, bytes]:
    """Export Markdown-ish text to PDF with proper CJK wrapping and basic styles.

    Uses ReportLab Platypus to provide automatic line wrapping, pagination, and
    list indentation. Japanese font is configured via built-in CID font.
    """
    buffer = BytesIO()

    # Register Japanese-capable CID font (built-in CJK font mapping)
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
    except Exception:
        # Registration is idempotent; ignore if already registered or unavailable
        pass

    # Document setup
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=40,
        rightMargin=40,
        topMargin=40,
        bottomMargin=40,
    )

    # Base styles with CJK wrapping
    base = ParagraphStyle(
        name="Base",
        fontName="HeiseiKakuGo-W5",
        fontSize=10,
        leading=14,
        spaceAfter=6,
        wordWrap="CJK",
    )
    style_title = ParagraphStyle(
        name="DocTitle",
        parent=base,
        fontSize=16,
        leading=20,
        spaceAfter=12,
    )
    style_h1 = ParagraphStyle(name="H1", parent=base, fontSize=14, leading=18, spaceBefore=8)
    style_h2 = ParagraphStyle(name="H2", parent=base, fontSize=12, leading=16, spaceBefore=6)
    style_h3 = ParagraphStyle(name="H3", parent=base, fontSize=11, leading=15, spaceBefore=4)
    style_quote = ParagraphStyle(
        name="Quote",
        parent=base,
        leftIndent=14,
        textColor="#444444",
        spaceBefore=4,
        spaceAfter=6,
    )
    style_code = ParagraphStyle(
        name="Code",
        parent=base,
        backColor="#F2F2F2",
        spaceBefore=4,
        spaceAfter=6,
        leading=13,
    )

    story: List[object] = []
    story.append(Paragraph(title or "特許明細書草案", style_title))

    # Inline Markdown -> simple XHTML for Paragraph
    def _md_inline_to_xhtml(text: str) -> str:
        # Split by backticks to protect code spans
        parts: list[str] = []
        i = 0
        while i < len(text):
            if text[i] == "`":
                j = text.find("`", i + 1)
                if j != -1:
                    code = text[i + 1 : j]
                    esc = code.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    parts.append(f"<font face='Courier'>{esc}</font>")
                    i = j + 1
                    continue
            # Non-code chunk
            j = text.find("`", i)
            segment = text[i : (j if j != -1 else len(text))]
            esc = segment.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            # Bold then italic (simple, non-nested)
            esc = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", esc)
            esc = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", esc)
            parts.append(esc)
            if j == -1:
                break
            i = j
        return "".join(parts)

    # Helper to flush a pending list buffer to story
    def flush_list(buffer: list[MarkdownElement]) -> None:
        if not buffer:
            return
        # Determine list type from first element
        first = buffer[0]
        items: list[ListItem] = []
        if first.kind == "bullet":
            for el in buffer:
                items.append(ListItem(Paragraph(_md_inline_to_xhtml(el.text), base)))
            story.append(
                ListFlowable(
                    items,
                    bulletType="bullet",
                    start=None,
                    bulletFontName=base.fontName,
                    leftPadding=18,
                )
            )
        elif first.kind == "number":
            start_num = first.number or 1
            for el in buffer:
                items.append(ListItem(Paragraph(_md_inline_to_xhtml(el.text), base)))
            story.append(
                ListFlowable(
                    items,
                    bulletType="1",
                    start=start_num,
                    bulletFontName=base.fontName,
                    leftPadding=18,
                )
            )

    # Accumulate elements and build story
    list_buffer: list[MarkdownElement] = []
    for element in parse_markdown(markdown_text or ""):
        if element.kind in ("bullet", "number"):
            # If switching between list types, flush
            if list_buffer and list_buffer[0].kind != element.kind:
                flush_list(list_buffer)
                list_buffer = []
            list_buffer.append(element)
            continue

        # Non-list element: flush any pending list
        flush_list(list_buffer)
        list_buffer = []

        if element.kind == "blank":
            story.append(Spacer(1, 8))
            continue

        if element.kind == "heading":
            level = element.level or 1
            if level <= 1:
                story.append(Paragraph(element.text, style_h1))
            elif level == 2:
                story.append(Paragraph(element.text, style_h2))
            else:
                story.append(Paragraph(element.text, style_h3))
            continue

        if element.kind == "quote":
            story.append(Paragraph(_md_inline_to_xhtml(element.text), style_quote))
            continue

        if element.kind == "codeblock":
            # Use a Paragraph per line to keep CJK font while showing a code-like block
            for ln in (element.text or "").splitlines() or [""]:
                story.append(Paragraph(_md_inline_to_xhtml(ln), style_code))
            continue

        # paragraph with inline styles
        story.append(Paragraph(_md_inline_to_xhtml(element.text), base))

    # Flush any trailing list
    flush_list(list_buffer)

    # Build PDF
    doc.build(story)
    return (f"{title or 'draft'}.pdf", buffer.getvalue())

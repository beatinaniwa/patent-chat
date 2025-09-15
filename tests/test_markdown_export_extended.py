from __future__ import annotations

from io import BytesIO

from docx import Document
from pypdf import PdfReader

from app.export import export_docx, export_pdf


def test_docx_inline_bold_italic_code_and_quote_and_codeblock():
    md = (
        "# タイトル\n\n"
        "> 引用の一文です\n\n"
        "本文の**太字**と*斜体*と`code`を含む段落。\n\n"
        "```\nprint('hello')\nコード行2\n```\n"
    )
    name, data = export_docx("タイトル", md)
    assert name.endswith(".docx")
    doc = Document(BytesIO(data))
    # Heading present
    assert any(p.style.name.startswith("Heading") and "タイトル" in p.text for p in doc.paragraphs)
    # Quote styled paragraph exists
    assert any("引用の一文です" in p.text for p in doc.paragraphs)

    # Find paragraph containing inline styles
    target = next(p for p in doc.paragraphs if "本文の" in p.text)
    texts = [
        (r.text, r.bold, r.italic, getattr(r.font, "highlight_color", None)) for r in target.runs
    ]
    assert any(t == "太字" and b for t, b, i, h in texts)
    assert any(t == "斜体" and i for t, b, i, h in texts)
    assert any(t == "code" for t, b, i, h in texts)

    # Code block lines present as separate paragraphs
    assert any("print('hello')" in p.text for p in doc.paragraphs)
    assert any("コード行2" in p.text for p in doc.paragraphs)


def test_pdf_extended_markdown_visible_text():
    md = (
        "# タイトル\n\n"
        "> 引用の一文です\n\n"
        "本文の**太字**と*斜体*と`code`を含む段落。\n\n"
        "```\nprint('hello')\nコード行2\n```\n"
    )
    name, data = export_pdf("タイトル", md)
    assert name.endswith(".pdf")
    reader = PdfReader(BytesIO(data))
    full_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    # Basic presence checks (styling not checkable via text extraction)
    for needle in [
        "タイトル",
        "引用の一文です",
        "太字",
        "斜体",
        "code",
        "print('hello')",
        "コード行2",
    ]:
        assert needle in full_text

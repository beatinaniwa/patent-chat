from __future__ import annotations

from io import BytesIO

from docx import Document
from pypdf import PdfReader

from app.export import export_docx, export_pdf


def test_markdown_rendered_in_docx_and_pdf():
    md = "# 見出し\n\n- 箇条書き\n\n通常の文章"
    docx_name, docx_bytes = export_docx("タイトル", md)
    assert docx_name.endswith(".docx")
    doc = Document(BytesIO(docx_bytes))
    texts = [p.text for p in doc.paragraphs]
    # Heading and bullet markers should be rendered, not literal
    assert "# 見出し" not in texts
    assert any(p.style.name.startswith("Heading") and p.text == "見出し" for p in doc.paragraphs)
    assert any(p.style.name == "List Bullet" and p.text == "箇条書き" for p in doc.paragraphs)

    pdf_name, pdf_bytes = export_pdf("タイトル", md)
    assert pdf_name.endswith(".pdf")
    reader = PdfReader(BytesIO(pdf_bytes))
    full_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "# 見出し" not in full_text
    assert "- 箇条書き" not in full_text
    assert "見出し" in full_text
    assert "箇条書き" in full_text
